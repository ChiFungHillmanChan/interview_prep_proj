"""
Document Parser Service

Handles PDF and DOCX file parsing with reliable text extraction.
Enforces file type and size validation.
"""

import re
import io
from typing import Optional, Tuple
from django.core.files.uploadedfile import UploadedFile
from django.core.exceptions import ValidationError

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None


class DocumentParserError(Exception):
    """Custom exception for document parsing errors"""
    pass


class DocumentParser:
    """
    Service class for parsing PDF and DOCX files to extract plain text.
    """
    
    # File size limits (in bytes)
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    # Allowed file types
    ALLOWED_EXTENSIONS = {'.pdf', '.docx'}
    ALLOWED_MIME_TYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }

    @classmethod
    def validate_file(cls, uploaded_file: UploadedFile) -> None:
        """
        Validate uploaded file type and size.
        
        Args:
            uploaded_file: Django uploaded file object
            
        Raises:
            ValidationError: If file is invalid
        """
        if not uploaded_file:
            raise ValidationError("No file provided")
        
        # Check file size
        if uploaded_file.size > cls.MAX_FILE_SIZE:
            max_mb = cls.MAX_FILE_SIZE / (1024 * 1024)
            raise ValidationError(f"File size exceeds {max_mb}MB limit")
        
        # Check file extension
        file_name = uploaded_file.name.lower()
        file_ext = None
        for ext in cls.ALLOWED_EXTENSIONS:
            if file_name.endswith(ext):
                file_ext = ext
                break
        
        if not file_ext:
            raise ValidationError(
                f"Invalid file type. Only PDF and DOCX files are allowed. "
                f"Received: {uploaded_file.name}"
            )
        
        # Check MIME type if available
        if hasattr(uploaded_file, 'content_type') and uploaded_file.content_type:
            if uploaded_file.content_type not in cls.ALLOWED_MIME_TYPES:
                raise ValidationError(
                    f"Invalid MIME type: {uploaded_file.content_type}. "
                    f"Only PDF and DOCX files are allowed."
                )

    @classmethod
    def parse_pdf(cls, file_content: bytes) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text string
            
        Raises:
            DocumentParserError: If parsing fails
        """
        if PyPDF2 is None:
            raise DocumentParserError(
                "PyPDF2 library not installed. Please install it to parse PDF files."
            )
        
        try:
            pdf_stream = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            if len(pdf_reader.pages) == 0:
                raise DocumentParserError("PDF file contains no pages")
            
            text_parts = []
            for page in pdf_reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    # Continue with other pages if one fails
                    continue
            
            if not text_parts:
                raise DocumentParserError("Could not extract text from PDF")
            
            return "\n\n".join(text_parts)
            
        except PyPDF2.errors.PdfReadError as e:
            raise DocumentParserError(f"Invalid PDF file: {e}")
        except Exception as e:
            raise DocumentParserError(f"Failed to parse PDF: {e}")

    @classmethod
    def parse_docx(cls, file_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text string
            
        Raises:
            DocumentParserError: If parsing fails
        """
        if Document is None:
            raise DocumentParserError(
                "python-docx library not installed. Please install it to parse DOCX files."
            )
        
        try:
            docx_stream = io.BytesIO(file_content)
            doc = Document(docx_stream)
            
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                raise DocumentParserError("Could not extract text from DOCX")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            raise DocumentParserError(f"Failed to parse DOCX: {e}")

    @classmethod
    def normalize_text(cls, text: str, max_length: int = 50000) -> str:
        """
        Normalize extracted text by cleaning whitespace and enforcing limits.
        
        Args:
            text: Raw extracted text
            max_length: Maximum text length
            
        Returns:
            Normalized text string
        """
        if not text:
            return ""
        
        # Normalize each line without erasing section boundaries used by the
        # deterministic Career Memory fallback.
        lines = [re.sub(r'[\t\r\f\v ]+', ' ', line).strip() for line in text.split('\n')]
        text = '\n'.join(lines).strip()
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Enforce length limit
        if len(text) > max_length:
            text = text[:max_length] + "..."
        
        return text

    @classmethod
    def parse_file(cls, uploaded_file: UploadedFile) -> Tuple[str, str]:
        """
        Parse uploaded file and extract text.
        
        Args:
            uploaded_file: Django uploaded file object
            
        Returns:
            Tuple of (extracted_text, file_type)
            
        Raises:
            ValidationError: If file validation fails
            DocumentParserError: If parsing fails
        """
        # Validate file first
        cls.validate_file(uploaded_file)
        
        # Read file content
        try:
            file_content = uploaded_file.read()
            uploaded_file.seek(0)  # Reset file pointer
        except Exception as e:
            raise DocumentParserError(f"Failed to read file: {e}")
        
        # Determine file type and parse
        file_name = uploaded_file.name.lower()
        
        if file_name.endswith('.pdf'):
            if not file_content.startswith(b'%PDF'):
                raise ValidationError('The uploaded file does not contain a valid PDF signature')
            raw_text = cls.parse_pdf(file_content)
            file_type = 'PDF'
        elif file_name.endswith('.docx'):
            if not file_content.startswith(b'PK'):
                raise ValidationError('The uploaded file does not contain a valid DOCX signature')
            raw_text = cls.parse_docx(file_content)
            file_type = 'DOCX'
        else:
            raise ValidationError("Unsupported file type")
        
        # Normalize the extracted text
        normalized_text = cls.normalize_text(raw_text)
        
        if not normalized_text.strip():
            raise DocumentParserError(f"No text could be extracted from the {file_type} file")
        
        return normalized_text, file_type

    @classmethod
    def get_preview_text(cls, text: str, max_words: int = 200) -> str:
        """
        Get a preview of the extracted text for display.
        
        Args:
            text: Full extracted text
            max_words: Maximum words in preview
            
        Returns:
            Preview text with ellipsis if truncated
        """
        if not text:
            return ""
        
        words = text.split()
        if len(words) <= max_words:
            return text
        
        preview_words = words[:max_words]
        return " ".join(preview_words) + "..."
