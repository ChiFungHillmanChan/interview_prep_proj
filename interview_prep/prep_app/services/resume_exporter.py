"""
Resume Export Service

Handles PDF and DOCX export functionality with A4 formatting.
Maintains exact layout and styling from the live preview.
"""

import os
import io
from typing import Dict, Any, Tuple
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.conf import settings
try:
    from ..schemas.ai_resume_schemas import EditableResume
except ImportError:
    from prep_app.schemas.ai_resume_schemas import EditableResume

try:
    import weasyprint
except ImportError:
    weasyprint = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    Document = None


class ResumeExporter:
    """
    Service class for exporting resumes to PDF and DOCX formats.
    """

    @classmethod
    def export_to_pdf(cls, resume_data: EditableResume, job_title: str = "") -> HttpResponse:
        """
        Export resume to PDF format using WeasyPrint.
        
        Args:
            resume_data: EditableResume object with complete resume data
            job_title: Job title for filename
            
        Returns:
            HttpResponse with PDF file
            
        Raises:
            Exception: If PDF generation fails
        """
        if weasyprint is None:
            raise Exception("WeasyPrint not installed. Please install it to export PDF.")

        try:
            # Generate HTML content
            html_content = cls._generate_html_for_export(resume_data)
            
            # Create PDF
            pdf_file = weasyprint.HTML(string=html_content).write_pdf()
            
            # Generate filename
            filename = cls._generate_filename(resume_data.name, job_title, 'pdf')
            
            # Create response
            response = HttpResponse(pdf_file, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except Exception as e:
            raise Exception(f"Failed to generate PDF: {str(e)}")

    @classmethod
    def export_to_docx(cls, resume_data: EditableResume, job_title: str = "") -> HttpResponse:
        """
        Export resume to DOCX format with styled document.
        
        Args:
            resume_data: EditableResume object with complete resume data
            job_title: Job title for filename
            
        Returns:
            HttpResponse with DOCX file
            
        Raises:
            Exception: If DOCX generation fails
        """
        if Document is None:
            raise Exception("python-docx not installed. Please install it to export DOCX.")

        try:
            # Create document
            doc = Document()
            
            # Configure document
            cls._configure_docx_document(doc)
            
            # Build resume content
            cls._build_docx_content(doc, resume_data)
            
            # Generate filename
            filename = cls._generate_filename(resume_data.name, job_title, 'docx')
            
            # Create response
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            response = HttpResponse(
                file_stream.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except Exception as e:
            raise Exception(f"Failed to generate DOCX: {str(e)}")

    @classmethod
    def _generate_html_for_export(cls, resume_data: EditableResume) -> str:
        """
        Generate HTML content optimized for PDF export.
        
        Args:
            resume_data: EditableResume object
            
        Returns:
            Complete HTML string with embedded CSS
        """
        # Get the export CSS (optimized for print)
        css_content = cls._get_export_css()
        
        # Render resume HTML
        resume_html = cls._render_resume_html(resume_data)
        
        # Combine into complete HTML document
        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Resume - {resume_data.name}</title>
    <style>
        {css_content}
    </style>
</head>
<body>
    <div class="ai-resume-page">
        {resume_html}
    </div>
</body>
</html>
"""
        return html_content

    @classmethod
    def _get_export_css(cls) -> str:
        """
        Get CSS optimized for PDF export.
        
        Returns:
            CSS string
        """
        return """
@page {
    size: A4;
    margin: 0;
}

body {
    margin: 0;
    padding: 0;
    font-family: "Times New Roman", Times, serif;
}

.ai-resume-page {
    box-sizing: border-box;
    width: 595.28pt;
    height: 841.89pt;
    padding: 28.2pt 15.5pt 32pt 21.5pt;
    background: #fff;
    page-break-after: always;
}

.ai-resume-name {
    font-family: "Times New Roman", Times, serif;
    font-weight: 700;
    font-size: 32pt;
    line-height: 1.05;
    text-align: center;
    margin: 0;
    color: #000;
}

.ai-resume-role {
    font-family: "Times New Roman", Times, serif;
    font-weight: 700;
    font-size: 19.5pt;
    text-align: center;
    margin: 0.6em 0 0;
    color: #000;
}

.ai-resume-contact-row {
    display: grid;
    grid-template-columns: 1fr auto 1fr;
    align-items: baseline;
    column-gap: 12pt;
    margin-top: 0.8em;
}

.ai-resume-contact-row .left,
.ai-resume-contact-row .center,
.ai-resume-contact-row .right {
    font-family: "Times New Roman", Times, serif;
    font-size: 11pt;
    line-height: 1.2;
    white-space: nowrap;
    min-width: 0;
    color: #000;
}

.ai-resume-contact-row .left { justify-self: start; }
.ai-resume-contact-row .center { justify-self: center; }
.ai-resume-contact-row .right { justify-self: end; text-align: right; }
.ai-resume-contact-row a { color: #000; text-decoration: none; }

.ai-resume-section { margin-top: 1.35em; }
.ai-resume-section-title,
.ai-resume-custom-section-title {
    font-family: "Times New Roman", Times, serif;
    font-weight: 700;
    font-size: 14pt;
    text-transform: uppercase;
    margin: 0;
    padding-bottom: 6pt;
    border-bottom: 1px solid #000;
    color: #000;
}

.ai-resume-summary {
    margin-top: 8pt;
    font-family: "Times New Roman", Times, serif;
    font-size: 14pt;
    line-height: 1.14;
    color: #000;
}

.ai-resume-skills-list { list-style: none; margin: 8pt 0 0; padding: 0; }
.ai-resume-skills-list li {
    font-family: "Times New Roman", Times, serif;
    font-size: 11pt;
    line-height: 1.2;
    color: #000;
}
.ai-resume-skills-list b { font-weight: 700; }

.ai-resume-edu-degree, .ai-resume-edu-institution {
    font-family: "Times New Roman", Times, serif;
    font-weight: 700;
    font-size: 14pt;
    margin: 0.6em 0 0;
    color: #000;
}

.ai-resume-edu-meta {
    display: flex;
    justify-content: space-between;
    font-family: "Times New Roman", Times, serif;
    font-size: 11pt;
    color: #000;
}

.ai-resume-entry-header {
    display: flex;
    justify-content: space-between;
    gap: 12pt;
    margin-top: 0.65em;
    align-items: flex-start;
}

.ai-resume-entry-title {
    font-family: "Times New Roman", Times, serif;
    font-weight: 700;
    font-size: 14pt;
    flex: 1 1 auto;
    min-width: 0;
    margin: 0;
    color: #000;
}

.ai-resume-entry-date {
    font-family: "Times New Roman", Times, serif;
    font-size: 11pt;
    white-space: nowrap;
    flex: 0 0 auto;
    margin: 0;
    color: #000;
}

.ai-resume-bullet-list { list-style: none; margin: 6pt 0 0; padding: 0; }
.ai-resume-bullet-list li {
    position: relative;
    font-family: "Times New Roman", Times, serif;
    font-size: 12pt;
    line-height: 1.22;
    margin: 0.25em 0 0;
    padding-left: 52.5pt;
    color: #000;
}
.ai-resume-bullet-list li::before {
    content: "•";
    position: absolute;
    left: 38.52pt;
    top: 0;
    font-size: 12pt;
}

.ai-resume-custom-section { margin-top: 1.35em; }
"""

    @classmethod
    def _render_resume_html(cls, resume_data: EditableResume) -> str:
        """
        Render resume data to HTML structure.
        
        Args:
            resume_data: EditableResume object
            
        Returns:
            HTML string for resume content
        """
        def escape_html(text):
            if not text:
                return ''
            return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        def render_skills_list(skills):
            categories = [
                ('Programming', skills.programming),
                ('Database', skills.database),
                ('AI/ML Tools', skills.ai_ml_tools),
                ('Tools & Methodologies', skills.tools_methodologies),
                ('Soft Skills', skills.soft_skills),
                ('Additional', skills.additional)
            ]
            
            items = []
            for label, skill_list in categories:
                if skill_list:
                    skill_text = ', '.join(skill_list) if isinstance(skill_list, list) else str(skill_list)
                    if skill_text.strip():
                        items.append(f'<li><b>{label}:</b> {escape_html(skill_text)}</li>')
            
            return ''.join(items)

        def render_education_list(education):
            if not education:
                return ''
            
            html_parts = []
            for edu in education:
                html_parts.append(f'<p class="ai-resume-edu-degree">{escape_html(edu.degree_bold)}</p>')
                html_parts.append(f'<p class="ai-resume-edu-institution">{escape_html(edu.institution_bold)}</p>')
                html_parts.append(f'''
                    <div class="ai-resume-edu-meta">
                        <div>{escape_html(edu.dates_left)}</div>
                        <div>{escape_html(edu.location_right)}</div>
                    </div>
                ''')
            
            return ''.join(html_parts)

        def render_entry_list(entries):
            if not entries:
                return ''
            
            html_parts = []
            for entry in entries:
                bullets = entry.bullets if isinstance(entry.bullets, list) else []
                bullet_html = ''
                if bullets:
                    bullet_items = ''.join([f'<li>{escape_html(bullet)}</li>' for bullet in bullets if bullet.strip()])
                    if bullet_items:
                        bullet_html = f'<ul class="ai-resume-bullet-list">{bullet_items}</ul>'
                
                html_parts.append(f'''
                    <div class="ai-resume-entry-header">
                        <p class="ai-resume-entry-title">{escape_html(entry.title_bold_left)}</p>
                        <p class="ai-resume-entry-date">{escape_html(entry.date_right_nowrap)}</p>
                    </div>
                    {bullet_html}
                ''')
            
            return ''.join(html_parts)

        def render_custom_sections(custom_sections):
            if not custom_sections:
                return ''
            
            html_parts = []
            for section in custom_sections:
                bullets = section.bullets if isinstance(section.bullets, list) else []
                bullet_html = ''
                if bullets:
                    bullet_items = ''.join([f'<li>{escape_html(bullet)}</li>' for bullet in bullets if bullet.strip()])
                    if bullet_items:
                        bullet_html = f'<ul class="ai-resume-bullet-list">{bullet_items}</ul>'
                
                html_parts.append(f'''
                    <section class="ai-resume-custom-section">
                        <h2 class="ai-resume-custom-section-title">{escape_html(section.heading)}</h2>
                        {bullet_html}
                    </section>
                ''')
            
            return ''.join(html_parts)

        # Build complete resume HTML
        html = f'''
            <h1 class="ai-resume-name">{escape_html(resume_data.name) if resume_data.name else 'Your Name'}</h1>
            <p class="ai-resume-role">{escape_html(resume_data.role) if resume_data.role else 'Your Role'}</p>
            
            <div class="ai-resume-contact-row">
                <div class="left">{escape_html(resume_data.contacts.email_left) if resume_data.contacts.email_left else 'email@example.com'}</div>
                <div class="center">{escape_html(resume_data.contacts.phone_center) if resume_data.contacts.phone_center else '+44 xxxx xxx xxx'}</div>
                <div class="right">{escape_html(resume_data.contacts.location_right) if resume_data.contacts.location_right else 'City, Country'}</div>
            </div>
            
            <div class="ai-resume-contact-row">
                <div class="left">GitHub: {f'<a href="{escape_html(resume_data.contacts.github_left)}">{escape_html(resume_data.contacts.github_left)}</a>' if resume_data.contacts.github_left else 'https://github.com/username'}</div>
                <div class="center"></div>
                <div class="right">Website: {f'<a href="{escape_html(resume_data.contacts.website_right)}">{escape_html(resume_data.contacts.website_right)}</a>' if resume_data.contacts.website_right else 'yourdomain.com'}</div>
            </div>

            <section class="ai-resume-section">
                <h2 class="ai-resume-section-title">SUMMARY</h2>
                <p class="ai-resume-summary">{escape_html(resume_data.summary) if resume_data.summary else 'Professional summary goes here.'}</p>
            </section>

            <section class="ai-resume-section">
                <h2 class="ai-resume-section-title">SKILLS</h2>
                <ul class="ai-resume-skills-list">
                    {render_skills_list(resume_data.skills)}
                </ul>
            </section>

            <section class="ai-resume-section">
                <h2 class="ai-resume-section-title">EDUCATION</h2>
                {render_education_list(resume_data.education)}
            </section>

            <section class="ai-resume-section">
                <h2 class="ai-resume-section-title">EXPERIENCE</h2>
                {render_entry_list(resume_data.experience)}
            </section>

            <section class="ai-resume-section">
                <h2 class="ai-resume-section-title">PROJECTS</h2>
                {render_entry_list(resume_data.projects)}
            </section>

            {render_custom_sections(resume_data.custom_sections)}
        '''
        
        return html

    @classmethod
    def _configure_docx_document(cls, doc: Document) -> None:
        """
        Configure DOCX document with proper margins and styles.
        
        Args:
            doc: python-docx Document object
        """
        # Set margins (A4 page margins to match template)
        section = doc.sections[0]
        section.top_margin = Inches(1.11)     # 28.2pt
        section.bottom_margin = Inches(1.26)  # 32pt
        section.left_margin = Inches(0.85)    # 21.5pt
        section.right_margin = Inches(0.61)   # 15.5pt
        
        # Create custom styles
        cls._create_docx_styles(doc)

    @classmethod
    def _create_docx_styles(cls, doc: Document) -> None:
        """
        Create custom styles for DOCX document.
        
        Args:
            doc: python-docx Document object
        """
        styles = doc.styles
        
        # Name style (32pt, bold, centered)
        name_style = styles.add_style('Resume Name', WD_STYLE_TYPE.PARAGRAPH)
        name_font = name_style.font
        name_font.name = 'Times New Roman'
        name_font.size = Pt(32)
        name_font.bold = True
        name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_style.paragraph_format.space_after = Pt(6)
        
        # Role style (19.5pt, bold, centered)
        role_style = styles.add_style('Resume Role', WD_STYLE_TYPE.PARAGRAPH)
        role_font = role_style.font
        role_font.name = 'Times New Roman'
        role_font.size = Pt(19.5)
        role_font.bold = True
        role_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role_style.paragraph_format.space_after = Pt(12)
        
        # Contact style (11pt)
        contact_style = styles.add_style('Resume Contact', WD_STYLE_TYPE.PARAGRAPH)
        contact_font = contact_style.font
        contact_font.name = 'Times New Roman'
        contact_font.size = Pt(11)
        contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_style.paragraph_format.space_after = Pt(4)
        
        # Section title style (14pt, bold, uppercase)
        section_style = styles.add_style('Resume Section', WD_STYLE_TYPE.PARAGRAPH)
        section_font = section_style.font
        section_font.name = 'Times New Roman'
        section_font.size = Pt(14)
        section_font.bold = True
        section_style.paragraph_format.space_before = Pt(18)
        section_style.paragraph_format.space_after = Pt(6)
        
        # Summary style (14pt)
        summary_style = styles.add_style('Resume Summary', WD_STYLE_TYPE.PARAGRAPH)
        summary_font = summary_style.font
        summary_font.name = 'Times New Roman'
        summary_font.size = Pt(14)
        summary_style.paragraph_format.space_after = Pt(8)
        
        # Entry title style (14pt, bold)
        entry_style = styles.add_style('Resume Entry', WD_STYLE_TYPE.PARAGRAPH)
        entry_font = entry_style.font
        entry_font.name = 'Times New Roman'
        entry_font.size = Pt(14)
        entry_font.bold = True
        entry_style.paragraph_format.space_before = Pt(8)
        entry_style.paragraph_format.space_after = Pt(2)
        
        # Bullet style (12pt)
        bullet_style = styles.add_style('Resume Bullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_font = bullet_style.font
        bullet_font.name = 'Times New Roman'
        bullet_font.size = Pt(12)
        bullet_style.paragraph_format.left_indent = Inches(0.5)
        bullet_style.paragraph_format.space_after = Pt(2)

    @classmethod
    def _build_docx_content(cls, doc: Document, resume_data: EditableResume) -> None:
        """
        Build DOCX content from resume data.
        
        Args:
            doc: python-docx Document object
            resume_data: EditableResume object
        """
        # Name and role
        name_para = doc.add_paragraph(resume_data.name or 'Your Name', style='Resume Name')
        role_para = doc.add_paragraph(resume_data.role or 'Your Role', style='Resume Role')
        
        # Contact information
        contact1 = doc.add_paragraph(style='Resume Contact')
        contact1.add_run(f"{resume_data.contacts.email_left or 'email@example.com'} • ")
        contact1.add_run(f"{resume_data.contacts.phone_center or '+44 xxxx xxx xxx'} • ")
        contact1.add_run(f"{resume_data.contacts.location_right or 'City, Country'}")
        
        contact2 = doc.add_paragraph(style='Resume Contact')
        if resume_data.contacts.github_left:
            contact2.add_run(f"GitHub: {resume_data.contacts.github_left}")
        if resume_data.contacts.website_right:
            if resume_data.contacts.github_left:
                contact2.add_run(" • ")
            contact2.add_run(f"Website: {resume_data.contacts.website_right}")
        
        # Summary section
        cls._add_docx_section(doc, "SUMMARY")
        if resume_data.summary:
            doc.add_paragraph(resume_data.summary, style='Resume Summary')
        
        # Skills section
        cls._add_docx_section(doc, "SKILLS")
        skills_categories = [
            ('Programming', resume_data.skills.programming),
            ('Database', resume_data.skills.database),
            ('AI/ML Tools', resume_data.skills.ai_ml_tools),
            ('Tools & Methodologies', resume_data.skills.tools_methodologies),
            ('Soft Skills', resume_data.skills.soft_skills),
            ('Additional', resume_data.skills.additional)
        ]
        
        for label, skills in skills_categories:
            if skills:
                skill_text = ', '.join(skills) if isinstance(skills, list) else str(skills)
                if skill_text.strip():
                    skill_para = doc.add_paragraph()
                    skill_para.add_run(f"{label}: ").bold = True
                    skill_para.add_run(skill_text)
        
        # Education section
        if resume_data.education:
            cls._add_docx_section(doc, "EDUCATION")
            for edu in resume_data.education:
                if edu.degree_bold:
                    edu_para = doc.add_paragraph(edu.degree_bold, style='Resume Entry')
                if edu.institution_bold:
                    inst_para = doc.add_paragraph(edu.institution_bold, style='Resume Entry')
                
                meta_para = doc.add_paragraph()
                if edu.dates_left:
                    meta_para.add_run(edu.dates_left)
                if edu.location_right:
                    if edu.dates_left:
                        # Create tab stop for right alignment simulation
                        meta_para.add_run(" " * 20)  # Approximate spacing
                    meta_para.add_run(edu.location_right)
        
        # Experience section
        if resume_data.experience:
            cls._add_docx_section(doc, "EXPERIENCE")
            cls._add_docx_entries(doc, resume_data.experience)
        
        # Projects section
        if resume_data.projects:
            cls._add_docx_section(doc, "PROJECTS")
            cls._add_docx_entries(doc, resume_data.projects)
        
        # Custom sections
        if resume_data.custom_sections:
            for section in resume_data.custom_sections:
                if section.heading and section.bullets:
                    cls._add_docx_section(doc, section.heading.upper())
                    for bullet in section.bullets:
                        if bullet.strip():
                            bullet_para = doc.add_paragraph(style='Resume Bullet')
                            bullet_para.add_run(f"• {bullet}")

    @classmethod
    def _add_docx_section(cls, doc: Document, title: str) -> None:
        """
        Add a section title to DOCX document.
        
        Args:
            doc: python-docx Document object
            title: Section title
        """
        section_para = doc.add_paragraph(title, style='Resume Section')
        # Add underline
        run = section_para.runs[0]
        run.underline = True

    @classmethod
    def _add_docx_entries(cls, doc: Document, entries: list) -> None:
        """
        Add experience/project entries to DOCX document.
        
        Args:
            doc: python-docx Document object
            entries: List of experience/project entries
        """
        for entry in entries:
            # Title and date on same line (approximate right alignment)
            title_para = doc.add_paragraph(style='Resume Entry')
            title_para.add_run(entry.title_bold_left or 'Title')
            if entry.date_right_nowrap:
                title_para.add_run(" " * 10)  # Approximate spacing
                date_run = title_para.add_run(entry.date_right_nowrap)
                date_run.bold = False
            
            # Bullets
            if entry.bullets:
                for bullet in entry.bullets:
                    if bullet.strip():
                        bullet_para = doc.add_paragraph(style='Resume Bullet')
                        bullet_para.add_run(f"• {bullet}")

    @classmethod
    def _generate_filename(cls, name: str, job_title: str, extension: str) -> str:
        """
        Generate appropriate filename for export.
        
        Args:
            name: Person's name
            job_title: Job title
            extension: File extension (pdf/docx)
            
        Returns:
            Generated filename
        """
        # Clean name
        clean_name = ''.join(c for c in (name or 'Resume') if c.isalnum() or c in ' -_').strip()
        clean_name = clean_name.replace(' ', '_')
        
        # Clean job title
        clean_job = ''.join(c for c in (job_title or '') if c.isalnum() or c in ' -_').strip()
        clean_job = clean_job.replace(' ', '_')
        
        # Build filename
        if clean_job:
            filename = f"{clean_name}_{clean_job}_Resume.{extension}"
        else:
            filename = f"{clean_name}_Resume.{extension}"
        
        # Ensure reasonable length
        if len(filename) > 100:
            filename = f"{clean_name[:30]}_Resume.{extension}"
        
        return filename