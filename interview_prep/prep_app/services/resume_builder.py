"""Truthful resume drafting against confirmed Career Memory."""

from __future__ import annotations

import io
import re
from html import escape
from typing import Any

from django.core.exceptions import ValidationError
from django.http import HttpResponse

from ..models import CareerMemoryFact
from .career_memory import memory_fingerprint


RESUME_SCHEMA_VERSION = 1
RESUME_SECTIONS = (
    'skills', 'experience', 'projects', 'education', 'achievements',
    'certifications', 'languages',
)


class TruthfulResumeService:
    """Build a canonical resume document; factual entries retain memory IDs."""

    @classmethod
    def build_document(cls, user, target_role: str, job_description: str) -> tuple[dict[str, Any], int]:
        facts = list(CareerMemoryFact.objects.filter(
            user=user,
            user_confirmed=True,
            review_status='confirmed',
        ).order_by('category', 'created_at'))
        by_category: dict[str, list[dict[str, Any]]] = {name: [] for name in RESUME_SECTIONS}
        personal = {'name': '', 'email': '', 'phone': '', 'location': '', 'website': '', 'github': ''}
        personal_memory_ids: dict[str, int] = {}

        category_to_section = {
            'skill': 'skills', 'work_experience': 'experience', 'experience': 'experience',
            'project': 'projects', 'education': 'education', 'achievement': 'achievements',
            'certification': 'certifications', 'language': 'languages',
        }
        for fact in facts:
            if fact.category == 'personal':
                field = str(fact.details.get('field', fact.title)).strip().casefold()
                aliases = {'full name': 'name', 'name': 'name', 'email': 'email', 'phone': 'phone',
                           'location': 'location', 'website': 'website', 'github': 'github'}
                if field in aliases:
                    canonical_field = aliases[field]
                    personal[canonical_field] = fact.content
                    personal_memory_ids[canonical_field] = fact.id
                continue
            section = category_to_section.get(fact.category)
            if section:
                by_category[section].append({
                    'memory_id': fact.id,
                    'title': fact.title,
                    'content': fact.content,
                    'details': fact.details,
                })

        confirmed_skill_names = [entry['content'] for entry in by_category['skills']]
        summary = ''
        if confirmed_skill_names:
            skills_text = ', '.join(confirmed_skill_names[:5])
            summary = f"Candidate targeting {target_role} roles with confirmed evidence in {skills_text}."
        elif target_role:
            summary = f"Candidate targeting {target_role} roles. Career evidence is still being confirmed."

        keywords = cls._job_keywords(job_description)
        evidence_text = ' '.join(f'{fact.title} {fact.content}' for fact in facts).casefold()
        matched = [word for word in keywords if word.casefold() in evidence_text]
        gaps = [word for word in keywords if word not in matched]
        coverage = round((len(matched) / len(keywords)) * 100) if keywords else 0
        questions = [
            f"Do you have real evidence for {gap}? If so, add and confirm it in Career Memory."
            for gap in gaps[:8]
        ]

        document = {
            'schema_version': RESUME_SCHEMA_VERSION,
            'personal': personal,
            'personal_memory_ids': personal_memory_ids,
            'target_role': target_role.strip()[:200],
            'summary': summary,
            **by_category,
            'job_match': {
                'coverage_percent': coverage,
                'matched_requirements': matched,
                'growth_areas': gaps,
                'questions': questions,
                'evidence_basis': f'{len(facts)} confirmed Career Memory item(s)',
            },
        }
        return document, coverage

    @staticmethod
    def _job_keywords(job_description: str) -> list[str]:
        if not job_description.strip():
            return []
        known_terms = (
            'Python', 'Django', 'Java', 'JavaScript', 'TypeScript', 'React', 'SQL',
            'PostgreSQL', 'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Git',
            'REST', 'GraphQL', 'Machine Learning', 'Data Analysis', 'Leadership',
            'Stakeholder Management', 'Communication', 'System Design', 'Agile',
        )
        return [term for term in known_terms if re.search(rf'\b{re.escape(term)}\b', job_description, re.I)]

    @classmethod
    def normalize_saved_document(cls, user, raw: Any) -> dict[str, Any]:
        if not isinstance(raw, dict) or raw.get('schema_version') != RESUME_SCHEMA_VERSION:
            raise ValidationError('Unsupported resume document schema')
        confirmed = {
            fact.id: fact for fact in CareerMemoryFact.objects.filter(
                user=user, user_confirmed=True, review_status='confirmed'
            )
        }
        cleaned = {
            'schema_version': RESUME_SCHEMA_VERSION,
            'personal': {},
            'target_role': str(raw.get('target_role', '')).strip()[:200],
            'summary': '',
        }
        personal = raw.get('personal') if isinstance(raw.get('personal'), dict) else {}
        personal_memory_ids = raw.get('personal_memory_ids') if isinstance(raw.get('personal_memory_ids'), dict) else {}
        cleaned['personal_memory_ids'] = {}
        for key in ('name', 'email', 'phone', 'location', 'website', 'github'):
            value = str(personal.get(key, '')).strip()[:250]
            if value:
                try:
                    memory_id = int(personal_memory_ids.get(key))
                except (TypeError, ValueError):
                    raise ValidationError('Personal details must come from confirmed Career Memory')
                fact = confirmed.get(memory_id)
                if not fact or fact.category != 'personal':
                    raise ValidationError('A personal detail is not backed by available confirmed evidence')
                if value != fact.content:
                    fingerprint = memory_fingerprint(fact.category, fact.title, value)
                    if CareerMemoryFact.objects.filter(user=user, fingerprint=fingerprint).exclude(id=fact.id).exists():
                        raise ValidationError('That edited personal detail already exists in Career Memory')
                    fact.content = value
                    fact.evidence = value
                    fact.source_type = 'manual'
                    fact.source_label = f'Edited in resume from {fact.source_label or fact.get_source_type_display()}'[:255]
                    fact.fingerprint = fingerprint
                    fact.save(update_fields=['content', 'evidence', 'source_type', 'source_label', 'fingerprint', 'updated_at'])
                cleaned['personal_memory_ids'][key] = fact.id
                cleaned['personal'][key] = fact.content
                # Included in the evidence basis even though personal fields are
                # rendered separately from resume sections.
            else:
                cleaned['personal'][key] = ''

        used_ids = set(cleaned['personal_memory_ids'].values())
        for section in RESUME_SECTIONS:
            cleaned[section] = []
            entries = raw.get(section) if isinstance(raw.get(section), list) else []
            for entry in entries[:50]:
                if not isinstance(entry, dict):
                    continue
                try:
                    memory_id = int(entry.get('memory_id'))
                except (TypeError, ValueError):
                    raise ValidationError('Every factual resume entry must come from confirmed Career Memory')
                fact = confirmed.get(memory_id)
                if not fact:
                    raise ValidationError('A resume entry refers to unavailable or unconfirmed evidence')
                # User edits made in the live editor become explicit, confirmed manual
                # evidence while keeping the original source excerpt for traceability.
                content = str(entry.get('content', fact.content)).strip()[:2000]
                title = str(entry.get('title', fact.title)).strip()[:250]
                if content != fact.content or title != fact.title:
                    fingerprint = memory_fingerprint(fact.category, title, content)
                    if CareerMemoryFact.objects.filter(user=user, fingerprint=fingerprint).exclude(id=fact.id).exists():
                        raise ValidationError('That edited resume evidence already exists in Career Memory')
                    fact.content = content
                    fact.title = title
                    fact.evidence = content
                    fact.source_type = 'manual'
                    fact.source_label = f'Edited in resume from {fact.source_label or fact.get_source_type_display()}'[:255]
                    fact.fingerprint = fingerprint
                    fact.save(update_fields=[
                        'content', 'title', 'evidence', 'source_type', 'source_label',
                        'fingerprint', 'updated_at'
                    ])
                cleaned[section].append({
                    'memory_id': fact.id,
                    'title': fact.title,
                    'content': fact.content,
                    'details': fact.details,
                })
                used_ids.add(fact.id)

        skill_names = [entry['content'] for entry in cleaned['skills']]
        if skill_names:
            cleaned['summary'] = (
                f"Candidate targeting {cleaned['target_role']} roles with confirmed evidence in "
                f"{', '.join(skill_names[:5])}."
            )
        elif cleaned['target_role']:
            cleaned['summary'] = (
                f"Candidate targeting {cleaned['target_role']} roles. Career evidence is still being confirmed."
            )

        job_match = raw.get('job_match') if isinstance(raw.get('job_match'), dict) else {}
        cleaned['job_match'] = {
            'coverage_percent': min(100, max(0, int(job_match.get('coverage_percent', 0) or 0))),
            'matched_requirements': cls._string_list(job_match.get('matched_requirements')),
            'growth_areas': cls._string_list(job_match.get('growth_areas')),
            'questions': cls._string_list(job_match.get('questions')),
            'evidence_basis': f'{len(used_ids)} confirmed Career Memory item(s)',
        }
        return cleaned

    @staticmethod
    def _string_list(value: Any) -> list[str]:
        if not isinstance(value, list):
            return []
        return list(dict.fromkeys(str(item).strip()[:250] for item in value if str(item).strip()))[:30]


class CareerResumeExporter:
    @classmethod
    def pdf_response(cls, document: dict[str, Any], filename: str) -> HttpResponse:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as exc:  # pragma: no cover - dependency check catches this
            raise RuntimeError('PDF export dependency is unavailable') from exc
        output = io.BytesIO()
        styles = getSampleStyleSheet()
        styles['Title'].textColor = colors.HexColor('#1558b0')
        pdf = SimpleDocTemplate(
            output, pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm,
            topMargin=18 * mm, bottomMargin=18 * mm,
        )
        story = []
        personal = document['personal']
        story.append(Paragraph(escape(personal.get('name') or 'Resume'), styles['Title']))
        if document.get('target_role'):
            story.append(Paragraph(escape(document['target_role']), styles['Heading2']))
        contacts = ' · '.join(value for value in personal.values() if value)
        if contacts:
            story.append(Paragraph(escape(contacts), styles['BodyText']))
        if document.get('summary'):
            story.extend([Spacer(1, 8), Paragraph('Profile', styles['Heading2']),
                          Paragraph(escape(document['summary']), styles['BodyText'])])
        for section in RESUME_SECTIONS:
            entries = document.get(section, [])
            if not entries:
                continue
            story.extend([Spacer(1, 8), Paragraph(section.replace('_', ' ').title(), styles['Heading2'])])
            for entry in entries:
                if entry.get('title'):
                    story.append(Paragraph(escape(entry['title']), styles['Heading3']))
                story.append(Paragraph(f"• {escape(entry.get('content', ''))}", styles['BodyText']))
        pdf.build(story)
        response = HttpResponse(output.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        return response

    @classmethod
    def docx_response(cls, document: dict[str, Any], filename: str) -> HttpResponse:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError('DOCX export dependency is unavailable') from exc
        output = io.BytesIO()
        doc = Document()
        personal = document['personal']
        doc.add_heading(personal.get('name') or 'Resume', 0)
        if document.get('target_role'):
            doc.add_paragraph(document['target_role'])
        contacts = ' · '.join(value for value in personal.values() if value)
        if contacts:
            doc.add_paragraph(contacts)
        if document.get('summary'):
            doc.add_heading('Profile', level=1)
            doc.add_paragraph(document['summary'])
        for section in RESUME_SECTIONS:
            entries = document.get(section, [])
            if not entries:
                continue
            doc.add_heading(section.replace('_', ' ').title(), level=1)
            for entry in entries:
                if entry.get('title'):
                    doc.add_paragraph(entry['title'], style='Heading 2')
                doc.add_paragraph(entry.get('content', ''), style='List Bullet')
        doc.save(output)
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response

    @classmethod
    def _html(cls, document: dict[str, Any]) -> str:
        personal = document['personal']
        contacts = ' · '.join(escape(value) for value in personal.values() if value)
        parts = [
            '<html><head><meta charset="utf-8"><style>',
            '@page{size:A4;margin:18mm}body{font:10.5pt Arial;color:#172033}',
            'h1{color:#1558b0;margin:0}h2{font-size:14pt;border-bottom:1px solid #bdd4f3;padding-bottom:3px}',
            'h3{font-size:11pt;margin-bottom:2px}p{line-height:1.4;margin:4px 0}',
            '</style></head><body>',
            f'<h1>{escape(personal.get("name") or "Resume")}</h1>',
            f'<p>{escape(document.get("target_role", ""))}</p><p>{contacts}</p>',
        ]
        if document.get('summary'):
            parts.extend(['<h2>Profile</h2>', f'<p>{escape(document["summary"])}</p>'])
        for section in RESUME_SECTIONS:
            entries = document.get(section, [])
            if entries:
                parts.append(f'<h2>{escape(section.replace("_", " ").title())}</h2>')
            for entry in entries:
                if entry.get('title'):
                    parts.append(f'<h3>{escape(entry["title"])}</h3>')
                parts.append(f'<p>• {escape(entry.get("content", ""))}</p>')
        parts.append('</body></html>')
        return ''.join(parts)
